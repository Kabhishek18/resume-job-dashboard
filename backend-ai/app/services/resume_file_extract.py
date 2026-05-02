"""Extract plain text from uploaded resume files (PDF, DOCX, TXT). Stored profile uses text only."""

from __future__ import annotations

import io
import re
from pathlib import PurePath

from app.core.errors import AppError


MAX_UPLOAD_BYTES = 10 * 1024 * 1024

_ALLOWED_EXT = {".pdf", ".docx", ".txt"}


def _normalize_ext(filename: str) -> str:
    return PurePath(filename or "").suffix.lower()


def extract_resume_plain_text(content: bytes, filename: str) -> tuple[str, list[str]]:
    """Return normalized plain_text and informational warnings."""
    warnings: list[str] = []

    if not content:
        raise AppError(
            "EMPTY_FILE",
            "Uploaded file was empty.",
            status_code=400,
        )

    if len(content) > MAX_UPLOAD_BYTES:
        raise AppError(
            "FILE_TOO_LARGE",
            f"File exceeds maximum size ({MAX_UPLOAD_BYTES // (1024 * 1024)} MB).",
            status_code=400,
        )

    ext = _normalize_ext(filename)
    if ext not in _ALLOWED_EXT:
        raise AppError(
            "UNSUPPORTED_TYPE",
            "Only PDF, DOCX, and TXT uploads are supported.",
            status_code=400,
        )

    if ext == ".txt":
        text = content.decode("utf-8", errors="replace").strip()
    elif ext == ".docx":
        text = _extract_docx(content)
        if len(text.strip()) < 40:
            warnings.append(
                "DOCX contained very little text — check formatting or paste manually if parsing looks wrong."
            )
    else:
        text = _extract_pdf(content)
        if len(text.strip()) < 40:
            warnings.append(
                "PDF may be image-only or scanned; little selectable text was found. Paste text if preview is blank."
            )

    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if len(text) > 500_000:
        text = text[:500_000]
        warnings.append("Extracted text was truncated at 500,000 characters.")

    if not text:
        raise AppError(
            "NO_TEXT_EXTRACTED",
            "Could not extract usable text from this file.",
            status_code=400,
        )

    return text, warnings


def _extract_docx(blob: bytes) -> str:
    from docx import Document

    fp = io.BytesIO(blob)
    doc = Document(fp)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_txt = "\t".join(c.text.strip() for c in row.cells).strip()
            if row_txt:
                parts.append(row_txt)
    return "\n".join(parts)


def _extract_pdf(blob: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise AppError(
            "PDF_UNAVAILABLE",
            "PDF libraries are unavailable on this server.",
            status_code=500,
        ) from e

    reader = PdfReader(io.BytesIO(blob))
    fragments: list[str] = []
    for page in reader.pages:
        try:
            fragments.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    return "\n\n".join(frag.strip() for frag in fragments if frag.strip())
