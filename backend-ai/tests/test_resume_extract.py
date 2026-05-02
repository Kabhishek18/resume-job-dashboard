import io
import uuid
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services import resume_file_extract as rfe


def test_extract_requires_auth():
    c = TestClient(app)
    r = c.post("/api/profile/resume-upload", files={"file": ("a.txt", b"hello resume", "text/plain")})
    assert r.status_code == 401


def _register_token() -> str:
    email = f"up-{uuid.uuid4().hex[:12]}@example.com"
    c = TestClient(app)
    r = c.post(
        "/api/auth/register",
        json={
            "name": "Up Test",
            "email": email,
            "password": "securepass123",
            "confirm_password": "securepass123",
        },
    )
    assert r.status_code == 200, r.json()
    return str(r.json()["access_token"])


def test_extract_txt_success():
    token = _register_token()
    c = TestClient(app)
    body = "Hello resume line one\nSenior engineer\nSecond line paragraph."
    r = c.post(
        "/api/profile/resume-upload",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("cv.txt", body.encode(), "text/plain")},
    )
    assert r.status_code == 200
    data = r.json()
    assert "Hello resume" in data["plain_text"]
    assert data["version"] == "v1"


def test_extract_docx_success():
    bio = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Docx headline role")
    doc.add_paragraph("Python FastAPI experience " * 5)
    doc.save(bio)
    blob = bio.getvalue()

    token = _register_token()
    c = TestClient(app)
    r = c.post(
        "/api/profile/resume-upload",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("cv.docx", blob, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    assert "FastAPI" in r.json()["plain_text"]


def test_extract_pdf_uses_extractor():
    token = _register_token()
    c = TestClient(app)
    fake_pdf = b"%PDF-1.4 placeholder bytes not a real pdf for reader"
    with patch.object(rfe, "_extract_pdf", return_value="Mocked selectable PDF lines " * 5):
        r = c.post(
            "/api/profile/resume-upload",
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("cv.pdf", fake_pdf, "application/pdf")},
        )
    assert r.status_code == 200
    assert "Mocked selectable PDF" in r.json()["plain_text"]
